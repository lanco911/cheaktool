"""
程序名称: CreateRecordDocument
程序功能: 生成检查记录表并保存为Word文档
作者: [您的名字]
日期: 2024年6月6日

概要:
本程序使用python-docx库生成一个标准化的检查记录表。用户可以指定任务ID、生产部门和文件保存路径，
程序将根据这些参数生成包含特定格式和内容的Word文档。生成的文档包括标题、初始表格、主表格、备注和签名部分。

主要功能:
1. 初始化类CreateRecordDocument并设置任务ID、生产部门和文件路径。
2. set_cell_border方法: 设置表格单元格的边框，包括边框宽度、颜色和样式。
3. set_font_style方法: 设置表格单元格内文本的字体样式，包括字体名称和大小。
4. set_column_width方法: 设置表格列的宽度。
5. create_table方法: 创建并保存检查记录表，包括添加标题、初始表格、主表格、备注和签名部分，
   并设置相应的字体样式和单元格边框。

使用方法:
1. 实例化CreateRecordDocument类，并传入任务ID、生产部门和文件保存路径。
2. 调用create_table方法生成并保存文档。

示例:
if __name__ == "__main__":
    task_id = "123456"  # 任务ID
    generation_department = "生产一部"  # 生产部门
    file_path = "检查记录表.docx"  # 保存文件路径
    document_creator = CreateRecordDocument(task_id, generation_department, file_path)
    document_creator.create_table()

依赖:
- python-docx: 一个用于创建和更新Microsoft Word (.docx) 文件的Python库。

注意事项:
- 确保已安装python-docx库，可以使用以下命令进行安装:
  pip install python-docx
"""

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from datetime import datetime

class CreateRecordDocument:
    def __init__(self, task_id, generation_department, file_path):
        """
        初始化方法，设置任务ID、生产部门和文件路径
        """
        self.task_id = task_id
        self.generation_department = generation_department
        self.file_path = file_path

    @staticmethod
    def set_cell_border(cell, **kwargs):
        """
        设置单元格边框
        :param cell: 需要设置边框的单元格
        :param kwargs: 边框参数，如 top={"sz": 12, "color": "#FF0000", "val": "single"}
        """
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        for border_name in ["top", "left", "bottom", "right"]:
            if border_name in kwargs:
                border = kwargs[border_name]
                element = tcPr.find(qn(f"w:{border_name}"))
                if element is None:
                    element = OxmlElement(f"w:{border_name}")
                    tcPr.append(element)
                if "sz" in border:
                    element.set(qn("w:sz"), str(border["sz"]))
                if "val" in border:
                    element.set(qn("w:val"), border["val"])
                if "color" in border:
                    element.set(qn("w:color"), border["color"])

    @staticmethod
    def set_font_style(cell, font_name="Microsoft Yahei", font_size=10):
        """
        设置单元格字体样式
        :param cell: 需要设置字体的单元格
        :param font_name: 字体名称，默认是“Microsoft Yahei”
        :param font_size: 字体大小，默认是10号
        """
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    @staticmethod
    def set_column_width(column, width):
        """
        设置表格列宽
        :param column: 需要设置宽度的列
        :param width: 列宽度，单位为厘米
        """
        for cell in column.cells:
            cell.width = width

    def create_table(self):
        """
        创建并保存检查记录表
        """
        # 创建一个新的Document对象
        doc = Document()

        # 设置页面边距：左右边距均为1.5厘米
        sections = doc.sections
        for section in sections:
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)

        # 添加标题段落
        p = doc.add_paragraph('检查记录表')

        # 设置标题字体样式
        for run in p.runs:
            run.font.name = 'Microsoft Yahei'
            run.font.size = Pt(18)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft Yahei')
        # 设置标题段落对齐方式为居中
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 添加初始表格，包含三行六列
        table = doc.add_table(rows=3, cols=6)

        # 填充第一行并合并单元格
        row = table.rows[0]
        row.cells[0].merge(row.cells[2])
        row.cells[3].merge(row.cells[5])
        row.cells[0].text = f"流 水 号：{self.task_id}"
        row.cells[3].text = f"生产部门：{self.generation_department}"

        # 填充第二行并合并所有单元格
        row = table.rows[1]
        row.cells[0].merge(row.cells[5])
        row.cells[0].text = "成果名称："

        # 填充第三行并合并所有单元格
        row = table.rows[2]
        row.cells[0].merge(row.cells[5])
        row.cells[0].text = "✓ ︎详查   ☐ 概查  ☐ 一级检查   ✓ ︎二级检查"
        # 设置段落对齐方式为右对齐
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # 设置初始表格字体样式
        for row in table.rows:
            for cell in row.cells:
                self.set_font_style(cell)

        # 创建主表格
        main_table = doc.add_table(rows=2, cols=11)
        # 设置各列宽度
        widths = [Cm(1.2), Cm(3), Cm(5.8), Cm(2), Cm(1.2), Cm(1.2), Cm(1.2), Cm(1.2), Cm(2.6), Cm(2.6), Cm(2.6)]
        for i, width in enumerate(widths):
            self.set_column_width(main_table.columns[i], width)

        # 合并并填充主表格表头行
        hdr_cells = main_table.rows[0].cells
        hdr_cells[0].text = '序号'
        hdr_cells[1].text = '编号'
        hdr_cells[2].text = '质量问题'
        hdr_cells[3].text = '质量元素或\n检查项'
        hdr_cells[4].merge(hdr_cells[7])
        hdr_cells[4].text = '错漏统计（个）'
        hdr_cells[8].text = '处理意见'
        hdr_cells[9].text = '修改情况'
        hdr_cells[10].text = '复核情况'

        # 填充并合并主表格的第二行
        sub_hdr_cells = main_table.rows[1].cells
        sub_hdr_cells[4].text = 'A类'
        sub_hdr_cells[5].text = 'B类'
        sub_hdr_cells[6].text = 'C类'
        sub_hdr_cells[7].text = 'D类'

        # 设置表头单元格段落对齐方式为居中
        for cell in hdr_cells:
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 设置表格单元格边框
        for row in main_table.rows:
            for cell in row.cells:
                self.set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                                          bottom={"sz": 12, "val": "single", "color": "000000"},
                                          left={"sz": 12, "val": "single", "color": "000000"},
                                          right={"sz": 12, "val": "single", "color": "000000"})

        # 合并主表格的部分列
        main_table.cell(0, 0).merge(main_table.cell(1, 0))
        main_table.cell(0, 1).merge(main_table.cell(1, 1))
        main_table.cell(0, 2).merge(main_table.cell(1, 2))
        main_table.cell(0, 3).merge(main_table.cell(1, 3))
        main_table.cell(0, 8).merge(main_table.cell(1, 8))
        main_table.cell(0, 9).merge(main_table.cell(1, 9))
        main_table.cell(0, 10).merge(main_table.cell(1, 10))

        # 添加10行空白行到主表格
        for _ in range(10):
            main_table.add_row()

        # 设置主表格字体样式
        for row in main_table.rows:
            for cell in row.cells:
                self.set_font_style(cell)

        # 添加备注表格
        footer_table = doc.add_table(rows=1, cols=1)
        footer_table.cell(0, 0).text = ("备注：\n"
                                        "1.质量元素：完整性填写“①”，属性精度填写“②”，要素及整饰质量填写“③”。\n"
                                        "2.质量评分：单位成果总分100分，A类错漏扣42分/个，B类错漏扣12分/个，C类错漏扣4分/个，D类错漏扣1分/个。\n"
                                        "3.检查结果：当符合率≥95%时为合格，反之则为不合格。")

        # 设置备注表格字体样式
        for row in footer_table.rows:
            for cell in row.cells:
                self.set_font_style(cell)

        # 获取当前日期并添加到签名表格中
        current_date = datetime.now().strftime("%Y-%m-%d")
        sign_off_table = doc.add_table(rows=1, cols=3)
        sign_off_table.cell(0, 0).text = f"检查者： 黄文勇\n检查日期：{current_date}"
        sign_off_table.cell(0, 1).text = "修改者：\n修改日期："
        sign_off_table.cell(0, 2).text = "复核者：\n复核日期："

        # 设置签名表格字体样式
        for row in sign_off_table.rows:
            for cell in row.cells:
                self.set_font_style(cell)

        # 设置所有表格的单元格边框
        for table in [table, main_table, footer_table, sign_off_table]:
            for row in table.rows:
                for cell in row.cells:
                    self.set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                                              bottom={"sz": 12, "val": "single", "color": "000000"},
                                              left={"sz": 12, "val": "single", "color": "000000"},
                                              right={"sz": 12, "val": "single", "color": "000000"})

        # 保存生成的文档
        doc.save(self.file_path)

if __name__ == "__main__":
    task_id = "123456"  # 任务ID
    generation_department = "生产一部"  # 生产部门
    file_path = "检查记录表.docx"  # 保存文件路径
    document_creator = CreateRecordDocument(task_id, generation_department, file_path)
    document_creator.create_table()
